
##Program Information
##Developer Information: Fly
##Developer Institution: JiangNan University (In WuXi, Jiangsu Provience, China)
##Contact me: COTOMO on GitHub
##Buy me a cup of coffee: Thanks, but no need currently


import docx
import os
import getpass
from os import system
import win32com.client

def save_doc_to_docx(path):  # docתdocx
    # ����WordӦ�ó������
    word_app = win32com.client.Dispatch("Word.Application")

    try:
        # ��doc�ļ�
        doc = word_app.Documents.Open(path)

        # ��doc�ļ����Ϊdocx
        doc.SaveAs2(path + "x", FileFormat=16)  # 16 ��ʾdocx��ʽ

        # �ر�doc�ļ�
        doc.Close()

    except Exception as e:
        print("�ڽ�docת��Ϊdocxʱ��������: {e}")

    finally:
        # �˳�WordӦ�ó���
        word_app.Quit()

class DocOperator:
    
    @staticmethod
    def main(doc):
        for paragraph in doc.paragraphs:
            text = paragraph.text
            start_spaces = len(text) - len(text.lstrip(' '))
            end_spaces = len(text) - len(text.rstrip(' '))
            total_length = len(text)

            if start_spaces > 0 and start_spaces < 6:  # ��������
                DocOperator.Indent(paragraph)
            if start_spaces > 6 and (total_length - start_spaces - end_spaces) <= 15: # �����Ű�
                DocOperator.Center(paragraph)    # �ļ��������Ӧ�����ⲿ����

    @staticmethod
    def Indent(paragraph):
        paragraph.text = paragraph.text.lstrip(' ')
        # ������������
        paragraph.paragraph_format.first_line_indent = docx.shared.Pt(24)  # ʾ����24��

    @staticmethod
    def Center(paragraph):
        paragraph.text = paragraph.text.strip()
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

class Informations:
    @staticmethod
    def Hello():
        print("\33[31mȥ��Ŀո��Ű淨��\33[0m")
        usrnm = getpass.getuser()
        print("��ã�\33[36m" + usrnm + "\33[0m!")
        print("��Ŀǰ����,�˳�����Խ����ĵ��е�ʹ�ÿո�������������;��е��Ű水�淶�ķ�ʽ����--��Ȼ,��Ҳֻ������ô�ࡣ")
        print("�˳�����\33[36mCOTOMO\33[0m������ά����")
        # ������Ϣ...

    @staticmethod
    def IsDocx(path):
        if path.lower().endswith(".docx"):
            return 0
        elif path.lower().endswith(".doc"):
            print("����ֻ�ܴ���docx�ļ�������doc�ļ�����docx�ļ���")
            save_doc_to_docx(path)
            return 1
        else:
            print("�������ֻ�ܴ���docx�ļ���doc�ļ��������ǻ������ġ�")
            return 2
        
class Configure:
    def developing():
        print("�˹������ڿ�����...")
    def main():
        CurrentSelection = 0
        Text = {"1.��������\n", "2. �����Ű�\n", "3. doc�ڴ�����ɺ��ٻ�ԭ��׺\n", "4.����Ϊ���ļ�\n", "5.�˳�\n"}
        Description = {"ѡ������Ƿ�ᴦ��������������","ѡ������Ƿ�ᴦ�����ľ����Ű�","ѡ������Ƿ���ڴ�����ɺ��ٻ�ԭ��׺","ѡ������Ƿ�ᱣ��Ϊ���ļ�","�˳�����"}
        while(1):
            print("����ҳ��")
            print("�������ɵ����ô˳������Ϊ���ұ�֤���޺��������ʹ��W���ƶ���꣬S���ƶ���꣬�Լ�ʹ��C�����Ĵ�ѡ�������")
            #�г����е�ѡ��,��������ǰѡ�е�ѡ��
            for i in range(len(Text)):
                if i == CurrentSelection:
                    print("\33[36m" + Text[i] + "\33[0m")
                else:
                    print(Text[i])
            print(' ')
            print(' ')
            print(Description[CurrentSelection])
            #��ȡ�û�����
            key = input()
            #�����û�����


#��Ҫ����
Informations.Hello()
FilePath = input("�����Ҫ�������ļ��Ͻ�����Ȼ��Enter����").strip('"')
if FilePath == "":
# �������������Ҫ����
    Configure.developing
elif Informations.IsDocx(FilePath) == 1:
    FilePath = FilePath + "x"
elif Informations.IsDocx(FilePath) == 2:
    print("�����޷�������ļ���")
    system("pause")

doc = docx.Document(FilePath) # ���ļ�
print(f"\33[31m�ɹ��ش����ļ� {FilePath} \33[0m���ļ������� \33[37m{len(doc.paragraphs)}\33[0m ��Paragraph")
print("��ʼ�����ļ�...")
DocOperator.main(doc)
# �����ļ�
doc.save(FilePath.strip(".docx") + "_�������.docx")
print("������ɣ��ļ���������" + FilePath)
print("�����������doc�ļ��Ļ���������һ��û�о���ԭʼ�Ű��޸Ĺ���docx�ļ�Ŷ��")
print("���Թص�����")
input()
